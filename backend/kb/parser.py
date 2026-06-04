from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document
from loguru import logger
from llmlingua import PromptCompressor


# ── LLMLingua compressor — initialized once at module load ────────
# Uses a small local model to compress text — no API call needed
# device_map="cpu" ensures it works without GPU
compressor = PromptCompressor(
    model_name="microsoft/llmlingua-2-xlm-roberta-large-meetingbank",
    use_llmlingua2=True,
    device_map="cpu",
)


def extract_text_from_file(file_path: str, content_type: str) -> str:
    """
    Main entry point — detects file type and routes to correct parser.
    Returns extracted and compressed plain text string.
    """
    try:
        if content_type == "application/pdf":
            return extract_from_pdf(file_path)

        elif content_type == (
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ):
            return extract_from_docx(file_path)

        else:
            logger.warning(f"Unsupported file type: {content_type}")
            return ""

    except Exception as e:
        logger.error(f"Text extraction failed for {file_path}: {e}")
        return ""


def extract_from_pdf(file_path: str) -> str:
    """
    Extracts text from a text-based PDF using pdfminer.six.
    Scanned image PDFs return empty string.
    """
    try:
        text = pdf_extract_text(file_path)

        if not text or not text.strip():
            logger.warning(
                f"No text extracted from PDF: {file_path}. "
                "PDF may be a scanned image."
            )
            return ""

        # Clean up excessive whitespace
        lines = [line.strip() for line in text.splitlines()]
        cleaned = "\n".join(line for line in lines if line)

        logger.info(
            f"PDF extracted: {file_path} | "
            f"Words: {len(cleaned.split())}"
        )

        # Compress if over 6000 words — preserve meaning, reduce tokens
        return compress_if_needed(cleaned, limit=6000)

    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return ""


def extract_from_docx(file_path: str) -> str:
    """
    Extracts text from a DOCX file using python-docx.
    Reads all paragraphs and table cells in order.
    """
    try:
        doc = Document(file_path)

        # Extract text from every paragraph
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables inside the document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in paragraphs:
                        paragraphs.append(cell_text)

        full_text = "\n".join(paragraphs)

        if not full_text.strip():
            logger.warning(f"No text extracted from DOCX: {file_path}")
            return ""

        logger.info(
            f"DOCX extracted: {file_path} | "
            f"Words: {len(full_text.split())}"
        )

        # Compress if over 6000 words — preserve meaning, reduce tokens
        return compress_if_needed(full_text, limit=6000)

    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return ""


def compress_if_needed(text: str, limit: int = 6000) -> str:
    """
    If text is within limit — return as is.
    If text exceeds limit — use LLMLingua to compress intelligently.
    LLMLingua preserves key information better than hard truncation.
    """
    words = text.split()

    # No compression needed
    if len(words) <= limit:
        logger.info(f"Text within limit ({len(words)} words) — no compression needed")
        return text

    logger.info(
        f"Text exceeds limit ({len(words)} words) — "
        f"compressing with LLMLingua..."
    )

    try:
        # Calculate compression ratio needed to hit target word count
        # target_token is approximate — LLMLingua works in tokens not words
        target_tokens = int(limit * 1.3)  # 1 word ≈ 1.3 tokens

        result = compressor.compress_prompt(
            text,
            # How aggressively to compress — 0.5 means keep 50% of tokens
            rate=limit / len(words),
            # Minimum tokens to keep regardless of rate
            target_token=target_tokens,
            # Preserve sentence structure for readability
            force_tokens=["\n", ".", "!", "?"],
        )

        compressed = result["compressed_prompt"]

        logger.info(
            f"LLMLingua compression complete | "
            f"Original: {len(words)} words → "
            f"Compressed: {len(compressed.split())} words"
        )

        return compressed

    except Exception as e:
        # If LLMLingua fails — fall back to hard truncation
        logger.warning(f"LLMLingua compression failed: {e} — falling back to truncation")
        return " ".join(words[:limit])


def get_word_count(text: str) -> int:
    """
    Returns word count of extracted text.
    Stored in DB and shown in KB panel.
    """
    if not text:
        return 0
    return len(text.split())